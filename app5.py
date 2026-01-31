import os
import numpy as np
import pandas as pd
import streamlit as st
from dotenv import load_dotenv
from docx import Document
from pypdf import PdfReader
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sentence_transformers import SentenceTransformer
import google.generativeai as genai





# .envファイルをロード
load_dotenv()

# APIキー取得
api_key = st.secrets["GOOGLE_API_KEY"]
genai.configure(api_key=api_key)
if not api_key:
    st.error("APIキーが設定されていません。Google CloudのAPIキーを設定してください。")
    st.stop()

# ================================
# 関数の設定
# ================================
# ================================
# Geminiモデル取得
# ================================
@st.cache_resource
def get_gemini_model():
    return genai.GenerativeModel("models/gemini-flash-latest")

# ================================
# Word / PDF 読み込み
# ================================
def load_documents_from_folder(folder_path):
    documents = []

    for file in os.listdir(folder_path):
        path = os.path.join(folder_path, file)

        if file.endswith(".docx"):
            content = load_word(path)
            documents.append({
                "content": content,
                "source": file,
                "type": "word",
                "location": "全文"
            })

        elif file.endswith(".pdf"):
            content = load_pdf(path)
            documents.append({
                "content": content,
                "source": file,
                "type": "pdf",
                "location": "全文"
            })

    return documents


def load_word(path):
    doc = Document(path)
    chunks = []

    for p in doc.paragraphs:
        if p.text.strip():
            chunks.append(p.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_text = " / ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                chunks.append(f"【表】{row_text}")

    return "\n".join(chunks)


def load_pdf(path):
    reader = PdfReader(path)
    chunks = []

    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text:
            text = text.replace("図", "【図】").replace("表", "【表】")
            chunks.append(f"【ページ{i+1}】\n{text}")

    return "\n".join(chunks)





# ================================
# CSV読み込み
# ================================
@st.cache_data
def load_data(csv_file_path):
    df = pd.read_csv(csv_file_path)
    df = df.dropna(subset=["documents"])
    return df


# ================================
# TF-IDFモデル構築
# ================================
@st.cache_resource
def build_tfidf_model(document):
    tfidf_vectorizer = TfidfVectorizer(
        max_features=5000,
        ngram_range=(1, 2),
        stop_words="english"
    )
    tfidf_matrix = tfidf_vectorizer.fit_transform(document)
    return tfidf_vectorizer, tfidf_matrix



# ================================
# SentenceTransformerモデル取得
# ================================
@st.cache_resource
def get_embedding_model():
    model = SentenceTransformer("all-MiniLM-L6-v2")
    return model


# ================================
# 埋め込みベクトル構築
# ================================
@st.cache_resource
def build_embedding_model(document):
    model = get_embedding_model()
    embeddings = model.encode(document, show_progress_bar=True)
    return embeddings


# ================================
# ハイブリッド検索
# ================================
def hybrid_search(query, tfidf_matrix, tfidf_vectorizer, embeddings, top_n=5):
    # TF-IDF類似度
    query_tfidf = tfidf_vectorizer.transform([query])
    tfidf_scores = cosine_similarity(query_tfidf, tfidf_matrix)[0]

    # Embedding類似度
    embed_model = get_embedding_model()
    query_embedding = embed_model.encode([query])
    embed_scores = cosine_similarity(query_embedding, embeddings)[0]

    # 正規化
    tfidf_scores = (tfidf_scores - tfidf_scores.min()) / (tfidf_scores.max() - tfidf_scores.min() + 1e-8)
    embed_scores = (embed_scores - embed_scores.min()) / (embed_scores.max() - embed_scores.min() + 1e-8)

    # ハイブリッドスコア
    hybrid_scores = (tfidf_scores + embed_scores) / 2

    top_indices = np.argsort(hybrid_scores)[::-1][:top_n]

    return top_indices, hybrid_scores[top_indices]



# ================================
# チャット履歴初期化
# ================================
def init_chat_history():
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []


# ================================
# チャット履歴表示
# ================================
def display_chat_history():
    for role, message in st.session_state.chat_history:
        with st.chat_message(role):
            st.markdown(message)


# ================================
# Geminiによる応答生成
# ================================
def respond_with_gemini(query, results, documents, top_n=3):
    model = get_gemini_model()

    context_parts = []
    for idx in results[:top_n]:
        d = documents[idx]
        context_parts.append(
            f"【出典】{d['source']}｜{d['location']}\n{d['content']}"
        )

    context = "\n\n".join(context_parts)

    prompt = f"""

あなたは【社内規程専用AI】です。
以下の【参照文書】に記載されている内容のみを根拠として回答してください。

【厳守事項】
- 図・表の内容も文章情報として解釈すること
- 判断できない場合は「参照文書に記載がありません」と明確に記載すること

【回答の整理ルール】
- 内容ごとに必ず見出し（##）を付ける
- 見出しは「結論が一目で分かる表現」にする
- 異なる話題を同一見出しに混在させない

【文章表現ルール】
- 長文を1段落にまとめない（2～3行以内で改行）
- 連続する文章に同一表現を多用しないこと
- わかりやすい口語表現にすること
- 単調な文章にならないように、文末を体言止めにしたり、Markdown形式を使用するなどして文章にリズム感を加えること

【禁止事項】
- 「〜と考えられる」「〜と思われる」などの曖昧表現
- 参照文書の記載を超えた言い換え


【参照文書】
{context}

【質問】
{query}

【回答】
"""

    response = model.generate_content(prompt)
    return response.text



# スライド作成
def generate_slide_markdown(query, results, documents, top_n=5):
    model = get_gemini_model()

    context_parts = []
    for idx in results[:top_n]:
        d = documents[idx]
        context_parts.append(
            f"【出典】{d['source']}｜{d['location']}\n{d['content']}"
        )

    context = "\n\n".join(context_parts)

    prompt = f"""
あなたは【社内資料スライド作成AI】です。
以下の【参照文書】のみを根拠に、スライド構成を作成してください。

【スライド作成ルール】
- 1スライド＝1メッセージ
- 推測・一般論は禁止
- 社内説明・QC発表向け
- 箇条書き中心
- 最大10枚まで

【出力形式（厳守）】
## スライドタイトル
- 要点1
- 要点2
- 要点3

## スライドタイトル
- 要点1
- 要点2

【参照文書】
{context}

【テーマ】
{query}

【スライド構成】
"""

    response = model.generate_content(prompt)
    return response.text



# wordに落とし込む関数
def slide_markdown_to_word(slide_md, out_path):
    doc = Document()

    lines = slide_md.splitlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # スライドタイトル → 見出し
        if line.startswith("## "):
            title = line.replace("## ", "").strip()
            doc.add_heading(title, level=1)

        # 箇条書き
        elif line.startswith("- "):
            bullet = line.replace("- ", "").strip()
            doc.add_paragraph(bullet, style="List Bullet")

        else:
            # 念のため通常文も受ける
            doc.add_paragraph(line)

    doc.save(out_path)




# ================================
# Streamlitアプリのメイン
# ================================
st.set_page_config(page_title="社内資料AI検索", layout="wide")

# ===== 社内向けUIスタイル =====
st.markdown("""
<style>
.chat-message-assistant {
    background-color: #f4f6f8;
}
.chat-message-user {
    background-color: #ffffff;
}
</style>
""", unsafe_allow_html=True)

st.title("📘 社内資料検索システム")

st.caption(
    "※ 本ツールは社内規程の検索補助を目的としています。"
    "最終判断は必ず原文をご確認ください。"
)


# ----------------
# フォルダから文書ロード
# ----------------
FOLDER_PATH = r"C:\Users\mt100\Downloads\テストAIチャットボット"


# ===== 検索用 corpus（documents と完全一致）=====
documents = load_documents_from_folder(FOLDER_PATH)

if not documents:
    st.error("文書がありません")
    st.stop()

corpus = [d["content"] for d in documents]

tfidf_vectorizer, tfidf_matrix = build_tfidf_model(corpus)
embeddings = build_embedding_model(corpus)




# ----------------
# チャット履歴初期化
# ----------------
init_chat_history()
display_chat_history()

# ----------------
# ユーザー入力
# ----------------
user_input = st.chat_input("質問を入力してください")

if user_input:
    # ユーザー発言を履歴に追加
    st.session_state.chat_history.append(("user", user_input))
    with st.chat_message("user"):
        st.markdown(user_input)
    
    # ----------------
    # 回答生成中表示
    # ----------------
    with st.spinner("🤖 回答を生成中です..."):
    
        # --------
        # 検索
        # --------
        top_indices, scores = hybrid_search(
        query=user_input,
        tfidf_matrix=tfidf_matrix,
        tfidf_vectorizer=tfidf_vectorizer,
        embeddings=embeddings,
        top_n=5
    )

        # 検索
        top_indices, scores = hybrid_search(
            query=user_input,
            tfidf_matrix=tfidf_matrix,
            tfidf_vectorizer=tfidf_vectorizer,
            embeddings=embeddings,
            top_n=5
        )

        valid_indices = list(top_indices)

        # ★ ここが重要
        st.session_state.valid_indices = valid_indices
        st.session_state.last_query = user_input


        # ★ しきい値なし：上位はすべて使う
        valid_indices = list(top_indices)

        # --------
        # Gemini応答生成
        # --------
        answer_body = respond_with_gemini(
            query=user_input,
            results=valid_indices,
            documents=documents,
            top_n=3
        )


        final_answer = answer_body + "\n\n" 

        st.session_state.chat_history.append(("assistant", final_answer))
        with st.chat_message("assistant"):
            st.markdown(final_answer)

    # --------
    # AI応答表示
    # --------
    with st.expander("🔍 参照した文書"):
        rows = []
        for idx in valid_indices[:3]:
            d = documents[idx]
            rows.append({
                "ファイル": d["source"],
                "種別": d["type"],
                "参照箇所": d["location"],
                "抜粋": d["content"][:200] + "…"
            })

        st.dataframe(pd.DataFrame(rows), use_container_width=True)


# ----------------
# スライド作成ボタン
# ----------------
st.markdown("### 📊 資料化")

if st.button("この質問をスライドにまとめる"):
    if "valid_indices" not in st.session_state:
        st.warning("先に質問を入力してください")
        st.stop()

    with st.spinner("📊 スライド構成を作成中..."):
        slide_md = generate_slide_markdown(
            query=st.session_state.last_query,
            results=st.session_state.valid_indices,
            documents=documents,
            top_n=5
        )

    # ★ session_state に保存
    st.session_state.slide_md = slide_md

# ----------------
# スライド構成案表示
# ----------------
if "slide_md" in st.session_state:
    st.markdown("## 🧾 スライド構成案")
    st.markdown(st.session_state.slide_md)

# ----------------
# Word
# ----------------
if st.button("📄 Wordに出力"):
    if "slide_md" not in st.session_state:
        st.warning("先にスライド構成案を作成してください")
        st.stop()

    os.makedirs("output", exist_ok=True)
    word_path = "output/スライド構成案.docx"

    slide_markdown_to_word(
        st.session_state.slide_md,
        word_path
    )

    st.download_button(
        label="📥 Wordをダウンロード",
        data=open(word_path, "rb"),
        file_name="スライド構成案.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )





